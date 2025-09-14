import os
import json
import sqlite3
import hashlib
import re
from typing import Dict, List, Optional, Any, Tuple, Union
from dataclasses import dataclass, field
from enum import Enum
from datetime import datetime, timedelta
from pathlib import Path
import tempfile
import base64

# Optional imports with fallback handling
try:
    from langchain_openai import ChatOpenAI, OpenAIEmbeddings
    from langchain.schema import Document
    from langchain.prompts import PromptTemplate

    LANGCHAIN_AVAILABLE = True
except ImportError:
    print("LangChain not available - using mock implementations")
    LANGCHAIN_AVAILABLE = False

# File processing imports
try:
    import PyPDF2
    import pdfplumber

    PDF_AVAILABLE = True
except ImportError:
    print("PDF processing libraries not available")
    PDF_AVAILABLE = False

try:
    from docx import Document as DocxDocument

    DOCX_AVAILABLE = True
except ImportError:
    print("DOCX processing library not available")
    DOCX_AVAILABLE = False

try:
    import pytesseract
    from PIL import Image

    OCR_AVAILABLE = True
except ImportError:
    print("OCR libraries not available")
    OCR_AVAILABLE = False

try:
    from deepeval import evaluate
    from deepeval.models import GPTModel
    from deepeval.metrics import GEval
    from deepeval.test_case import LLMTestCase

    DEEPEVAL_AVAILABLE = True
except ImportError:
    print("DeepEval not available - using mock evaluation")
    DEEPEVAL_AVAILABLE = False


# ========================================================================================
# CONFIGURATION AND DATA STRUCTURES
# ========================================================================================


class SkillCategory(Enum):
    """Skill classification enumeration"""

    TECHNICAL = "TECHNICAL"
    LEADERSHIP = "LEADERSHIP"
    COMMUNICATION = "COMMUNICATION"
    ANALYTICAL = "ANALYTICAL"
    CREATIVE = "CREATIVE"
    INTERPERSONAL = "INTERPERSONAL"
    LANGUAGE = "LANGUAGE"
    PROJECT_MANAGEMENT = "PROJECT_MANAGEMENT"


class ProficiencyLevel(Enum):
    """Proficiency level enumeration"""

    BEGINNER = 1
    INTERMEDIATE = 2
    ADVANCED = 3
    EXPERT = 4
    MASTER = 5


class CulturalDimension(Enum):
    """Cultural fit dimensions"""

    TEAMWORK = "TEAMWORK"
    INNOVATION = "INNOVATION"
    INTEGRITY = "INTEGRITY"
    ADAPTABILITY = "ADAPTABILITY"
    LEARNING_AGILITY = "LEARNING_AGILITY"
    COMMUNICATION_STYLE = "COMMUNICATION_STYLE"
    WORK_STYLE = "WORK_STYLE"


@dataclass
class ReaderConfig:
    """Resume reader configuration"""

    database_path: str = "resumes/resume_analysis.db"
    temp_dir: str = "temp_files"
    supported_formats: List[str] = field(
        default_factory=lambda: [".pdf", ".docx", ".txt", ".png", ".jpg", ".jpeg"]
    )
    enable_ocr: bool = True
    ocr_language: str = "eng"
    min_confidence_threshold: float = 0.6
    enable_cultural_analysis: bool = True
    skill_extraction_depth: str = "comprehensive"  # basic, standard, comprehensive


@dataclass
class PersonalInfo:
    """Personal basic information"""

    full_name: str
    email: Optional[str] = None
    phone: Optional[str] = None
    address: Optional[str] = None
    linkedin: Optional[str] = None
    github: Optional[str] = None
    portfolio: Optional[str] = None
    summary: Optional[str] = None


@dataclass
class Education:
    """Education background"""

    degree: str
    major: str
    school: str
    graduation_year: Optional[int] = None
    gpa: Optional[float] = None
    honors: List[str] = field(default_factory=list)
    relevant_courses: List[str] = field(default_factory=list)


@dataclass
class WorkExperience:
    """Work experience"""

    position: str
    company: str
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    duration_months: Optional[int] = None
    responsibilities: List[str] = field(default_factory=list)
    achievements: List[str] = field(default_factory=list)
    technologies_used: List[str] = field(default_factory=list)


@dataclass
class Skill:
    """Skill information"""

    name: str
    category: SkillCategory
    proficiency_level: ProficiencyLevel
    years_experience: Optional[int] = None
    proficiency_score: float = 0.0  # 0-100 scale
    evidence_sources: List[str] = field(default_factory=list)


@dataclass
class CulturalFitScore:
    """Cultural fit assessment score"""

    dimension: CulturalDimension
    score: float  # 0-100 scale
    evidence: List[str] = field(default_factory=list)
    reasoning: str = ""


@dataclass
class CompetencyAssessment:
    """Competency assessment"""

    competency_name: str
    current_score: float  # 0-100 scale
    potential_score: float  # potential assessment score
    evidence: List[str] = field(default_factory=list)
    development_areas: List[str] = field(default_factory=list)


@dataclass
class ResumeAnalysisResult:
    """Resume analysis result"""

    analysis_id: str
    personal_info: PersonalInfo
    education: List[Education]
    work_experience: List[WorkExperience]
    skills: List[Skill]
    cultural_fit_scores: List[CulturalFitScore]
    competency_assessments: List[CompetencyAssessment]
    overall_score: float  # overall score 0-100
    early_talent_suitability: float  # early talent suitability 0-100
    leadership_potential: float  # leadership potential 0-100
    career_progression_analysis: str
    strengths: List[str]
    improvement_areas: List[str]
    red_flags: List[str]
    analysis_timestamp: datetime
    file_hash: str
    confidence_score: float  # analysis accuracy 0-1.0
    processing_metadata: Dict[str, Any]


# ========================================================================================
# MOCK IMPLEMENTATIONS
# ========================================================================================


class MockLLM:
    """Mock LLM implementation"""

    def __init__(self, **kwargs):
        self.kwargs = kwargs

    def invoke(self, prompt: str) -> str:
        """Return mock responses based on prompt content"""
        prompt_lower = prompt.lower()

        if "extract personal info" in prompt_lower:
            return """
            NAME: John Smith
            EMAIL: john.smith@email.com
            PHONE: +1-555-0123
            SUMMARY: Experienced software developer with 2 years in web development
            """
        elif "extract skills" in prompt_lower:
            return """
            TECHNICAL: Python (Advanced), JavaScript (Intermediate), React (Intermediate)
            LEADERSHIP: Team Collaboration (Intermediate), Project Management (Beginner)
            COMMUNICATION: Written Communication (Advanced), Public Speaking (Intermediate)
            """
        elif "assess cultural fit" in prompt_lower:
            return """
            TEAMWORK: 75 | Strong collaboration skills evident
            INNOVATION: 65 | Shows creative problem-solving approach
            ADAPTABILITY: 80 | Demonstrates flexibility in different roles
            """
        elif "evaluate" in prompt_lower or "score" in prompt_lower:
            return "SCORE: 0.8 | EXPLANATION: Comprehensive analysis with good evidence extraction"
        else:
            return "Mock analysis result for resume processing"


# ========================================================================================
# FILE PROCESSORS
# ========================================================================================


class FileProcessor:
    """Unified file processor"""

    def __init__(self, config: ReaderConfig):
        self.config = config
        self.temp_dir = Path(config.temp_dir)
        self.temp_dir.mkdir(parents=True, exist_ok=True)

    def process_file(self, file_path: str) -> str:
        """Process file and return text content"""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        file_extension = file_path.suffix.lower()

        if file_extension == ".pdf":
            return self._process_pdf(file_path)
        elif file_extension == ".docx":
            return self._process_docx(file_path)
        elif file_extension == ".txt":
            return self._process_txt(file_path)
        elif file_extension in [".png", ".jpg", ".jpeg"]:
            return self._process_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")

    def _process_pdf(self, file_path: Path) -> str:
        """Process PDF file"""
        if not PDF_AVAILABLE:
            return f"PDF processing not available. File: {file_path.name}"

        try:
            # Try using pdfplumber (better table and layout handling)
            with pdfplumber.open(file_path) as pdf:
                text = ""
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"

                if text.strip():
                    return text

            # Fallback to PyPDF2
            with open(file_path, "rb") as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text

        except Exception as e:
            print(f"Error processing PDF {file_path}: {e}")
            return f"Error processing PDF file: {str(e)}"

    def _process_docx(self, file_path: Path) -> str:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            return f"DOCX processing not available. File: {file_path.name}"

        try:
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"

            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + "\t"
                    text += "\n"

            return text

        except Exception as e:
            print(f"Error processing DOCX {file_path}: {e}")
            return f"Error processing DOCX file: {str(e)}"

    def _process_txt(self, file_path: Path) -> str:
        """Process TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as file:
                return file.read()
        except UnicodeDecodeError:
            # Try other encodings
            encodings = ["gbk", "gb2312", "latin-1"]
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue

            return f"Error: Unable to decode text file {file_path}"
        except Exception as e:
            print(f"Error processing TXT {file_path}: {e}")
            return f"Error processing TXT file: {str(e)}"

    def _process_image(self, file_path: Path) -> str:
        """Process image file (OCR)"""
        if not OCR_AVAILABLE or not self.config.enable_ocr:
            return f"OCR processing not available or disabled. File: {file_path.name}"

        try:
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang=self.config.ocr_language)
            return text
        except Exception as e:
            print(f"Error processing image {file_path}: {e}")
            return f"Error processing image file: {str(e)}"


# ========================================================================================
# INFORMATION EXTRACTORS
# ========================================================================================


class InformationExtractor:
    """Information extractor"""

    def __init__(self, llm, config: ReaderConfig):
        self.llm = llm
        self.config = config

        # Information extraction prompt templates
        self.personal_info_prompt = (
            PromptTemplate(
                template="""
            Extract personal information from the following resume text.

            Resume text:
            {resume_text}

            Please extract information in the following format:
            NAME: [Full name]
            EMAIL: [Email address]
            PHONE: [Phone number]
            ADDRESS: [Address]
            LINKEDIN: [LinkedIn URL]
            GITHUB: [GitHub URL]  
            PORTFOLIO: [Portfolio URL]
            SUMMARY: [Personal summary/objective statement]

            If any information is not found, please mark as "N/A".
            """,
                input_variables=["resume_text"],
            )
            if LANGCHAIN_AVAILABLE
            else None
        )

        self.skills_extraction_prompt = (
            PromptTemplate(
                template="""
            Extract skill information from the following resume text, categorize by type and assess proficiency.

            Resume text:
            {resume_text}

            Please extract skills in the following format:
            TECHNICAL: [Technical skills] (Level: Beginner/Intermediate/Advanced/Expert/Master)
            LEADERSHIP: [Leadership skills] (Level)
            COMMUNICATION: [Communication skills] (Level)
            ANALYTICAL: [Analytical skills] (Level)
            CREATIVE: [Creative skills] (Level)
            INTERPERSONAL: [Interpersonal skills] (Level)
            LANGUAGE: [Language skills] (Level)
            PROJECT_MANAGEMENT: [Project management skills] (Level)

            Assess levels based on years of experience, project complexity, and job requirements.
            """,
                input_variables=["resume_text"],
            )
            if LANGCHAIN_AVAILABLE
            else None
        )

        self.cultural_fit_prompt = (
            PromptTemplate(
                template="""
            Analyze cultural fit based on resume content. Assess candidate performance in the following dimensions (0-100 scale):

            Resume text:
            {resume_text}

            Please assess in the following format:
            TEAMWORK: [Score] | [Evidence description]
            INNOVATION: [Score] | [Evidence description]
            INTEGRITY: [Score] | [Evidence description]
            ADAPTABILITY: [Score] | [Evidence description]
            LEARNING_AGILITY: [Score] | [Evidence description]
            COMMUNICATION_STYLE: [Score] | [Evidence description]
            WORK_STYLE: [Score] | [Evidence description]

            Provide evidence based on work history, project experience, education background, etc.
            """,
                input_variables=["resume_text"],
            )
            if LANGCHAIN_AVAILABLE
            else None
        )

    def extract_personal_info(self, resume_text: str) -> PersonalInfo:
        """Extract personal basic information"""
        try:
            if self.personal_info_prompt:
                prompt = self.personal_info_prompt.format(resume_text=resume_text)
            else:
                prompt = f"Extract personal information from resume:\n{resume_text}"

            response = self.llm.invoke(prompt)

            if hasattr(response, "content"):
                response_text = response.content
            else:
                response_text = str(response)

            return self._parse_personal_info(response_text, resume_text)

        except Exception as e:
            print(f"Error extracting personal info: {e}")
            return self._extract_personal_info_fallback(resume_text)

    def _parse_personal_info(self, llm_response: str, resume_text: str) -> PersonalInfo:
        """Parse LLM returned personal information"""
        info = {}
        lines = llm_response.strip().split("\n")

        for line in lines:
            if ":" in line:
                key, value = line.split(":", 1)
                key = key.strip().upper()
                value = value.strip()
                if value and value != "N/A":
                    info[key] = value

        # Fallback to regex extraction
        if not info.get("NAME"):
            name_match = re.search(r"^([A-Za-z\s]{2,50})", resume_text, re.MULTILINE)
            if name_match:
                info["NAME"] = name_match.group(1).strip()

        if not info.get("EMAIL"):
            email_match = re.search(
                r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", resume_text
            )
            if email_match:
                info["EMAIL"] = email_match.group()

        if not info.get("PHONE"):
            phone_match = re.search(r"[\+]?[1-9]?[\d\s\-\(\)]{10,15}", resume_text)
            if phone_match:
                info["PHONE"] = phone_match.group()

        return PersonalInfo(
            full_name=info.get("NAME", "Unknown"),
            email=info.get("EMAIL"),
            phone=info.get("PHONE"),
            address=info.get("ADDRESS"),
            linkedin=info.get("LINKEDIN"),
            github=info.get("GITHUB"),
            portfolio=info.get("PORTFOLIO"),
            summary=info.get("SUMMARY"),
        )

    def _extract_personal_info_fallback(self, resume_text: str) -> PersonalInfo:
        """Fallback method when personal info extraction fails"""
        # Use regex for basic extraction
        name_match = re.search(r"^([A-Za-z\s]{2,50})", resume_text, re.MULTILINE)
        email_match = re.search(
            r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b", resume_text
        )
        phone_match = re.search(r"[\+]?[1-9]?[\d\s\-\(\)]{10,15}", resume_text)

        return PersonalInfo(
            full_name=name_match.group(1).strip() if name_match else "Unknown",
            email=email_match.group() if email_match else None,
            phone=phone_match.group() if phone_match else None,
        )

    def extract_skills(self, resume_text: str) -> List[Skill]:
        """Extract skill information"""
        try:
            if self.skills_extraction_prompt:
                prompt = self.skills_extraction_prompt.format(resume_text=resume_text)
            else:
                prompt = f"Extract skills from resume:\n{resume_text}"

            response = self.llm.invoke(prompt)

            if hasattr(response, "content"):
                response_text = response.content
            else:
                response_text = str(response)

            return self._parse_skills(response_text)

        except Exception as e:
            print(f"Error extracting skills: {e}")
            return self._extract_skills_fallback(resume_text)

    def _parse_skills(self, llm_response: str) -> List[Skill]:
        """Parse skill information"""
        skills = []
        lines = llm_response.strip().split("\n")

        level_mapping = {
            "BEGINNER": ProficiencyLevel.BEGINNER,
            "INTERMEDIATE": ProficiencyLevel.INTERMEDIATE,
            "ADVANCED": ProficiencyLevel.ADVANCED,
            "EXPERT": ProficiencyLevel.EXPERT,
            "MASTER": ProficiencyLevel.MASTER,
        }

        for line in lines:
            if ":" in line:
                try:
                    category_part, skills_part = line.split(":", 1)
                    category_name = category_part.strip().upper()

                    if category_name in [c.value for c in SkillCategory]:
                        category = SkillCategory[category_name]

                        # Parse skills and levels
                        skill_entries = skills_part.split(",")
                        for skill_entry in skill_entries:
                            skill_entry = skill_entry.strip()
                            if skill_entry:
                                # Extract skill name and level
                                level_match = re.search(r"\((.*?)\)", skill_entry)
                                if level_match:
                                    level_text = level_match.group(1).strip().upper()
                                    skill_name = skill_entry.replace(
                                        level_match.group(0), ""
                                    ).strip()

                                    level = level_mapping.get(
                                        level_text, ProficiencyLevel.INTERMEDIATE
                                    )
                                    score = self._level_to_score(level)

                                    skills.append(
                                        Skill(
                                            name=skill_name,
                                            category=category,
                                            proficiency_level=level,
                                            proficiency_score=score,
                                        )
                                    )
                except Exception as e:
                    print(f"Error parsing skill line: {line}, error: {e}")
                    continue

        return skills

    def _level_to_score(self, level: ProficiencyLevel) -> float:
        """Convert level to score"""
        mapping = {
            ProficiencyLevel.BEGINNER: 20.0,
            ProficiencyLevel.INTERMEDIATE: 40.0,
            ProficiencyLevel.ADVANCED: 60.0,
            ProficiencyLevel.EXPERT: 80.0,
            ProficiencyLevel.MASTER: 95.0,
        }
        return mapping.get(level, 40.0)

    def _extract_skills_fallback(self, resume_text: str) -> List[Skill]:
        """Fallback method when skill extraction fails"""
        # Basic skill recognition based on keywords
        technical_keywords = [
            "python",
            "java",
            "javascript",
            "react",
            "sql",
            "git",
            "aws",
            "docker",
        ]
        soft_skills = ["leadership", "communication", "teamwork", "problem solving"]

        skills = []
        resume_lower = resume_text.lower()

        for keyword in technical_keywords:
            if keyword in resume_lower:
                skills.append(
                    Skill(
                        name=keyword.title(),
                        category=SkillCategory.TECHNICAL,
                        proficiency_level=ProficiencyLevel.INTERMEDIATE,
                        proficiency_score=50.0,
                    )
                )

        for keyword in soft_skills:
            if keyword in resume_lower:
                skills.append(
                    Skill(
                        name=keyword.title(),
                        category=SkillCategory.INTERPERSONAL,
                        proficiency_level=ProficiencyLevel.INTERMEDIATE,
                        proficiency_score=50.0,
                    )
                )

        return skills

    def extract_cultural_fit(self, resume_text: str) -> List[CulturalFitScore]:
        """Extract cultural fit analysis"""
        try:
            if self.cultural_fit_prompt:
                prompt = self.cultural_fit_prompt.format(resume_text=resume_text)
            else:
                prompt = f"Analyze cultural fit from resume:\n{resume_text}"

            response = self.llm.invoke(prompt)

            if hasattr(response, "content"):
                response_text = response.content
            else:
                response_text = str(response)

            return self._parse_cultural_fit(response_text)

        except Exception as e:
            print(f"Error extracting cultural fit: {e}")
            return self._extract_cultural_fit_fallback()

    def _parse_cultural_fit(self, llm_response: str) -> List[CulturalFitScore]:
        """Parse cultural fit analysis"""
        cultural_scores = []
        lines = llm_response.strip().split("\n")

        for line in lines:
            if ":" in line and "|" in line:
                try:
                    dimension_part, rest = line.split(":", 1)
                    dimension_name = dimension_part.strip().upper()

                    if dimension_name in [d.value for d in CulturalDimension]:
                        score_part, evidence_part = rest.split("|", 1)
                        score = float(score_part.strip())
                        evidence = evidence_part.strip()

                        cultural_scores.append(
                            CulturalFitScore(
                                dimension=CulturalDimension[dimension_name],
                                score=score,
                                evidence=[evidence],
                                reasoning=evidence,
                            )
                        )
                except Exception as e:
                    print(f"Error parsing cultural fit line: {line}, error: {e}")
                    continue

        return cultural_scores

    def _extract_cultural_fit_fallback(self) -> List[CulturalFitScore]:
        """Fallback method when cultural fit analysis fails"""
        # Return default medium scores
        default_scores = []
        for dimension in CulturalDimension:
            default_scores.append(
                CulturalFitScore(
                    dimension=dimension,
                    score=60.0,
                    evidence=["Analysis not available"],
                    reasoning="Default scoring due to analysis failure",
                )
            )
        return default_scores


# ========================================================================================
# COMPETENCY ASSESSOR
# ========================================================================================


class CompetencyAssessor:
    """Competency assessor"""

    def __init__(self, llm, config: ReaderConfig):
        self.llm = llm
        self.config = config

        # Mapping to existing LeadershipCompetency
        self.competency_mapping = {
            "COMMUNICATION": "Communication Skills",
            "EMOTIONAL_INTELLIGENCE": "Emotional Intelligence",
            "DECISION_MAKING": "Decision Making",
            "STRATEGIC_THINKING": "Strategic Thinking",
            "TEAM_BUILDING": "Team Building",
            "ADAPTABILITY": "Adaptability",
            "INFLUENCE": "Influence",
            "ACCOUNTABILITY": "Accountability",
        }

    def assess_competencies(
        self,
        resume_text: str,
        skills: List[Skill],
        work_experience: List[WorkExperience],
    ) -> List[CompetencyAssessment]:
        """Assess various competencies"""
        assessments = []

        for competency_key, competency_name in self.competency_mapping.items():
            assessment = self._assess_single_competency(
                competency_key, competency_name, resume_text, skills, work_experience
            )
            assessments.append(assessment)

        return assessments

    def _assess_single_competency(
        self,
        competency_key: str,
        competency_name: str,
        resume_text: str,
        skills: List[Skill],
        work_experience: List[WorkExperience],
    ) -> CompetencyAssessment:
        """Assess single competency"""

        prompt = f"""
        Based on the following resume information, assess the candidate's {competency_name} competency:

        Resume text:
        {resume_text}

        Please provide assessment in the following format:
        CURRENT_SCORE: [0-100 score]
        POTENTIAL_SCORE: [0-100 score] 
        EVIDENCE: [Specific evidence, separated by semicolons]
        DEVELOPMENT_AREAS: [Development suggestions, separated by semicolons]
        REASONING: [Scoring rationale]

        Scoring criteria:
        - 0-20: Beginner level, needs extensive guidance
        - 21-40: Basic level, can work with guidance
        - 41-60: Intermediate level, can handle common situations independently
        - 61-80: Advanced level, can guide others
        - 81-100: Expert level, can innovate and lead
        """

        try:
            response = self.llm.invoke(prompt)

            if hasattr(response, "content"):
                response_text = response.content
            else:
                response_text = str(response)

            return self._parse_competency_assessment(competency_name, response_text)

        except Exception as e:
            print(f"Error assessing competency {competency_name}: {e}")
            return self._create_default_assessment(competency_name)

    def _parse_competency_assessment(
        self, competency_name: str, llm_response: str
    ) -> CompetencyAssessment:
        """Parse competency assessment result"""

        lines = llm_response.strip().split("\n")
        parsed = {}

        for line in lines:
            if ":" in line:
                key, value = line.split(":", 1)
                parsed[key.strip().upper()] = value.strip()

        try:
            current_score = float(parsed.get("CURRENT_SCORE", "50"))
            potential_score = float(parsed.get("POTENTIAL_SCORE", "60"))
            evidence = (
                parsed.get("EVIDENCE", "").split(";") if parsed.get("EVIDENCE") else []
            )
            development_areas = (
                parsed.get("DEVELOPMENT_AREAS", "").split(";")
                if parsed.get("DEVELOPMENT_AREAS")
                else []
            )

            return CompetencyAssessment(
                competency_name=competency_name,
                current_score=max(0, min(100, current_score)),
                potential_score=max(0, min(100, potential_score)),
                evidence=[e.strip() for e in evidence if e.strip()],
                development_areas=[d.strip() for d in development_areas if d.strip()],
            )

        except Exception as e:
            print(f"Error parsing competency assessment: {e}")
            return self._create_default_assessment(competency_name)

    def _create_default_assessment(self, competency_name: str) -> CompetencyAssessment:
        """Create default assessment"""
        return CompetencyAssessment(
            competency_name=competency_name,
            current_score=50.0,
            potential_score=60.0,
            evidence=["Insufficient information for assessment"],
            development_areas=["Need more information for detailed assessment"],
        )


# ========================================================================================
# DATABASE MANAGER
# ========================================================================================


class DatabaseManager:
    """Database manager"""

    def __init__(self, config: ReaderConfig):
        self.config = config
        self.db_path = Path(config.database_path)
        self.db_path.parent.mkdir(parents=True, exist_ok=True)
        self._init_database()

    def _init_database(self):
        """Initialize database table structure"""
        with sqlite3.connect(self.db_path) as conn:
            cursor = conn.cursor()

            # Main resume analysis table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS resume_analyses (
                    analysis_id TEXT PRIMARY KEY,
                    full_name TEXT NOT NULL,
                    analysis_timestamp DATETIME NOT NULL,
                    file_hash TEXT,
                    overall_score REAL,
                    early_talent_suitability REAL,
                    leadership_potential REAL,
                    confidence_score REAL,
                    career_progression_analysis TEXT,
                    processing_metadata TEXT
                )
            """
            )

            # Personal info table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS personal_info (
                    analysis_id TEXT,
                    full_name TEXT,
                    email TEXT,
                    phone TEXT,
                    address TEXT,
                    linkedin TEXT,
                    github TEXT,
                    portfolio TEXT,
                    summary TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Education table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS education (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    degree TEXT,
                    major TEXT,
                    school TEXT,
                    graduation_year INTEGER,
                    gpa REAL,
                    honors TEXT,
                    relevant_courses TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Work experience table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS work_experience (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    position TEXT,
                    company TEXT,
                    start_date TEXT,
                    end_date TEXT,
                    duration_months INTEGER,
                    responsibilities TEXT,
                    achievements TEXT,
                    technologies_used TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Skills table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS skills (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    skill_name TEXT,
                    category TEXT,
                    proficiency_level INTEGER,
                    proficiency_score REAL,
                    years_experience INTEGER,
                    evidence_sources TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Cultural fit table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS cultural_fit (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    dimension TEXT,
                    score REAL,
                    evidence TEXT,
                    reasoning TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Competency assessments table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS competency_assessments (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    competency_name TEXT,
                    current_score REAL,
                    potential_score REAL,
                    evidence TEXT,
                    development_areas TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Analysis insights table
            cursor.execute(
                """
                CREATE TABLE IF NOT EXISTS analysis_insights (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    analysis_id TEXT,
                    insight_type TEXT,  -- 'strength', 'improvement_area', 'red_flag'
                    content TEXT,
                    FOREIGN KEY (analysis_id) REFERENCES resume_analyses (analysis_id)
                )
            """
            )

            # Create indexes
            cursor.execute(
                "CREATE INDEX IF NOT EXISTS idx_name_timestamp ON resume_analyses (full_name, analysis_timestamp)"
            )
            cursor.execute(
                "CREATE INDEX IF NOT EXISTS idx_analysis_id ON personal_info (analysis_id)"
            )

            conn.commit()

    def save_analysis_result(self, result: ResumeAnalysisResult) -> bool:
        """Save analysis result to database"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()

                # Save main analysis record
                cursor.execute(
                    """
                    INSERT OR REPLACE INTO resume_analyses 
                    (analysis_id, full_name, analysis_timestamp, file_hash, overall_score, 
                     early_talent_suitability, leadership_potential, confidence_score, 
                     career_progression_analysis, processing_metadata)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                    (
                        result.analysis_id,
                        result.personal_info.full_name,
                        result.analysis_timestamp.isoformat(),
                        result.file_hash,
                        result.overall_score,
                        result.early_talent_suitability,
                        result.leadership_potential,
                        result.confidence_score,
                        result.career_progression_analysis,
                        json.dumps(result.processing_metadata),
                    ),
                )

                # Save personal info
                self._save_personal_info(
                    cursor, result.analysis_id, result.personal_info
                )

                # Save education
                for edu in result.education:
                    self._save_education(cursor, result.analysis_id, edu)

                # Save work experience
                for exp in result.work_experience:
                    self._save_work_experience(cursor, result.analysis_id, exp)

                # Save skills
                for skill in result.skills:
                    self._save_skill(cursor, result.analysis_id, skill)

                # Save cultural fit
                for cultural_fit in result.cultural_fit_scores:
                    self._save_cultural_fit(cursor, result.analysis_id, cultural_fit)

                # Save competency assessments
                for assessment in result.competency_assessments:
                    self._save_competency_assessment(
                        cursor, result.analysis_id, assessment
                    )

                # Save analysis insights
                self._save_analysis_insights(
                    cursor,
                    result.analysis_id,
                    result.strengths,
                    result.improvement_areas,
                    result.red_flags,
                )

                conn.commit()
                return True

        except Exception as e:
            print(f"Error saving analysis result: {e}")
            return False

    def _save_personal_info(self, cursor, analysis_id: str, info: PersonalInfo):
        """Save personal information"""
        cursor.execute(
            """
            INSERT OR REPLACE INTO personal_info 
            (analysis_id, full_name, email, phone, address, linkedin, github, portfolio, summary)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                info.full_name,
                info.email,
                info.phone,
                info.address,
                info.linkedin,
                info.github,
                info.portfolio,
                info.summary,
            ),
        )

    def _save_education(self, cursor, analysis_id: str, edu: Education):
        """Save education"""
        cursor.execute(
            """
            INSERT INTO education 
            (analysis_id, degree, major, school, graduation_year, gpa, honors, relevant_courses)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                edu.degree,
                edu.major,
                edu.school,
                edu.graduation_year,
                edu.gpa,
                json.dumps(edu.honors),
                json.dumps(edu.relevant_courses),
            ),
        )

    def _save_work_experience(self, cursor, analysis_id: str, exp: WorkExperience):
        """Save work experience"""
        cursor.execute(
            """
            INSERT INTO work_experience 
            (analysis_id, position, company, start_date, end_date, duration_months, 
             responsibilities, achievements, technologies_used)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                exp.position,
                exp.company,
                exp.start_date,
                exp.end_date,
                exp.duration_months,
                json.dumps(exp.responsibilities),
                json.dumps(exp.achievements),
                json.dumps(exp.technologies_used),
            ),
        )

    def _save_skill(self, cursor, analysis_id: str, skill: Skill):
        """Save skill"""
        cursor.execute(
            """
            INSERT INTO skills 
            (analysis_id, skill_name, category, proficiency_level, proficiency_score, 
             years_experience, evidence_sources)
            VALUES (?, ?, ?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                skill.name,
                skill.category.value,
                skill.proficiency_level.value,
                skill.proficiency_score,
                skill.years_experience,
                json.dumps(skill.evidence_sources),
            ),
        )

    def _save_cultural_fit(
        self, cursor, analysis_id: str, cultural_fit: CulturalFitScore
    ):
        """Save cultural fit"""
        cursor.execute(
            """
            INSERT INTO cultural_fit 
            (analysis_id, dimension, score, evidence, reasoning)
            VALUES (?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                cultural_fit.dimension.value,
                cultural_fit.score,
                json.dumps(cultural_fit.evidence),
                cultural_fit.reasoning,
            ),
        )

    def _save_competency_assessment(
        self, cursor, analysis_id: str, assessment: CompetencyAssessment
    ):
        """Save competency assessment"""
        cursor.execute(
            """
            INSERT INTO competency_assessments 
            (analysis_id, competency_name, current_score, potential_score, evidence, development_areas)
            VALUES (?, ?, ?, ?, ?, ?)
        """,
            (
                analysis_id,
                assessment.competency_name,
                assessment.current_score,
                assessment.potential_score,
                json.dumps(assessment.evidence),
                json.dumps(assessment.development_areas),
            ),
        )

    def _save_analysis_insights(
        self,
        cursor,
        analysis_id: str,
        strengths: List[str],
        improvement_areas: List[str],
        red_flags: List[str],
    ):
        """Save analysis insights"""
        for strength in strengths:
            cursor.execute(
                """
                INSERT INTO analysis_insights (analysis_id, insight_type, content)
                VALUES (?, ?, ?)
            """,
                (analysis_id, "strength", strength),
            )

        for improvement in improvement_areas:
            cursor.execute(
                """
                INSERT INTO analysis_insights (analysis_id, insight_type, content)
                VALUES (?, ?, ?)
            """,
                (analysis_id, "improvement_area", improvement),
            )

        for red_flag in red_flags:
            cursor.execute(
                """
                INSERT INTO analysis_insights (analysis_id, insight_type, content)
                VALUES (?, ?, ?)
            """,
                (analysis_id, "red_flag", red_flag),
            )

    def get_candidate_history(self, full_name: str) -> List[Dict[str, Any]]:
        """Get candidate history analysis records"""
        try:
            with sqlite3.connect(self.db_path) as conn:
                cursor = conn.cursor()
                cursor.execute(
                    """
                    SELECT analysis_id, analysis_timestamp, overall_score, 
                           early_talent_suitability, leadership_potential, confidence_score
                    FROM resume_analyses 
                    WHERE full_name = ?
                    ORDER BY analysis_timestamp DESC
                """,
                    (full_name,),
                )

                results = cursor.fetchall()
                history = []

                for row in results:
                    history.append(
                        {
                            "analysis_id": row[0],
                            "timestamp": row[1],
                            "overall_score": row[2],
                            "early_talent_suitability": row[3],
                            "leadership_potential": row[4],
                            "confidence_score": row[5],
                        }
                    )

                return history

        except Exception as e:
            print(f"Error retrieving candidate history: {e}")
            return []


# ========================================================================================
# G-EVAL ASSESSOR
# ========================================================================================


class GEvalAssessor:
    """G-Eval assessor"""

    def __init__(self, config: ReaderConfig):
        self.config = config
        self.evaluator_model = None

        if DEEPEVAL_AVAILABLE and os.getenv("OPENAI_API_KEY"):
            try:
                self.evaluator_model = GPTModel(model="gpt-4o-mini")
                print("G-Eval evaluator initialized successfully.")
            except Exception as e:
                print(f"Failed to initialize G-Eval: {e}")

    def evaluate_extraction_quality(
        self, original_text: str, extracted_result: ResumeAnalysisResult
    ) -> float:
        """Evaluate resume extraction quality"""

        if not self.evaluator_model:
            return self._mock_evaluation(extracted_result)

        try:
            # Create G-Eval evaluation criteria
            evaluation_criteria = """
            Evaluate the quality of resume information extraction considering the following dimensions:
            1. Information Completeness (30%): Whether all important information was extracted
            2. Accuracy (25%): Whether extracted information is accurate and error-free  
            3. Structuring Level (20%): Whether information was properly classified and organized
            4. Competency Assessment Reasonableness (15%): Whether skill and competency scores are reasonable
            5. Analysis Depth (10%): Whether valuable insights were provided
            
            Score range: 0-1.0, where 1.0 indicates perfect extraction
            """

            # Prepare evaluation input
            extracted_summary = self._create_extraction_summary(extracted_result)

            test_case = LLMTestCase(
                input=f"Original resume text: {original_text[:2000]}...",  # Limit length
                actual_output=extracted_summary,
                expected_output="High-quality structured resume information extraction",
            )

            # Create G-Eval evaluation metric
            g_eval_metric = GEval(
                name="Resume Extraction Quality",
                criteria=evaluation_criteria,
                evaluation_params=[
                    "Information Completeness",
                    "Accuracy",
                    "Structuring Level",
                    "Competency Assessment Reasonableness",
                    "Analysis Depth",
                ],
                model=self.evaluator_model,
                threshold=0.6,
            )

            # Run evaluation
            evaluation_results = evaluate(
                test_cases=[test_case], metrics=[g_eval_metric]
            )

            # Extract score
            if evaluation_results.test_results:
                score = evaluation_results.test_results[0].metrics_data[0].score
                return float(score)

            return 0.7  # Default score

        except Exception as e:
            print(f"G-Eval evaluation error: {e}")
            return self._mock_evaluation(extracted_result)

    def _create_extraction_summary(self, result: ResumeAnalysisResult) -> str:
        """Create extraction result summary"""
        summary = f"""
        Extraction Result Summary:
        - Candidate: {result.personal_info.full_name}
        - Contact Info: Email {result.personal_info.email}, Phone {result.personal_info.phone}
        - Skills Count: {len(result.skills)}
        - Work Experience: {len(result.work_experience)} positions
        - Education: {len(result.education)} degrees
        - Overall Score: {result.overall_score}/100
        - Early Talent Suitability: {result.early_talent_suitability}/100
        - Leadership Potential: {result.leadership_potential}/100
        - Cultural Fit Dimensions: {len(result.cultural_fit_scores)}
        - Competency Assessments: {len(result.competency_assessments)}
        - Strengths Identified: {len(result.strengths)} items
        - Improvement Areas: {len(result.improvement_areas)} items
        """
        return summary

    def _mock_evaluation(self, result: ResumeAnalysisResult) -> float:
        """Mock evaluation"""
        # Simple scoring based on extraction completeness
        score = 0.5  # Base score

        # Basic information completeness
        if result.personal_info.email:
            score += 0.1
        if result.personal_info.phone:
            score += 0.1

        # Skills and experience
        if len(result.skills) > 3:
            score += 0.1
        if len(result.work_experience) > 0:
            score += 0.1

        # Assessment quality
        if result.overall_score > 0:
            score += 0.1

        return min(score, 1.0)


# ========================================================================================
# MAIN READER WORKFLOW
# ========================================================================================


class ReaderWorkflow:
    """Main resume reader workflow"""

    def __init__(self, llm, config: ReaderConfig = None):
        self.config = config or ReaderConfig()
        self.llm = llm
        self.initialization_time = datetime.now()

        try:
            print("Initializing resume reader system...")

            # Initialize components
            self.file_processor = FileProcessor(self.config)
            self.information_extractor = InformationExtractor(llm, self.config)
            self.competency_assessor = CompetencyAssessor(llm, self.config)
            self.database_manager = DatabaseManager(self.config)
            self.geval_assessor = GEvalAssessor(self.config)

            self.initialization_successful = True
            print("Resume reader system initialized successfully.")

        except Exception as e:
            print(f"Error initializing reader workflow: {e}")
            self.initialization_successful = False
            self.initialization_error = str(e)

    def process_resume(
        self, file_path: str, candidate_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """Process single resume file"""

        if not self.initialization_successful:
            return self._create_initialization_error_response()

        start_time = datetime.now()
        processing_steps = []

        try:
            # Step 1: Process file
            processing_steps.append("Processing file and extracting text...")
            resume_text = self.file_processor.process_file(file_path)

            if not resume_text or len(resume_text.strip()) < 50:
                return self._create_error_response(
                    "File content is empty or too short", processing_steps
                )

            # Calculate file hash
            file_hash = hashlib.md5(resume_text.encode()).hexdigest()

            # Step 2: Extract basic information
            processing_steps.append("Extracting personal information...")
            personal_info = self.information_extractor.extract_personal_info(
                resume_text
            )

            # Use provided candidate name if available
            if candidate_name:
                personal_info.full_name = candidate_name

            # Step 3: Extract skills
            processing_steps.append("Extracting skills and competencies...")
            skills = self.information_extractor.extract_skills(resume_text)

            # Step 4: Extract work experience (simplified implementation)
            processing_steps.append("Analyzing work experience...")
            work_experience = self._extract_work_experience_simple(resume_text)

            # Step 5: Extract education (simplified implementation)
            processing_steps.append("Extracting education background...")
            education = self._extract_education_simple(resume_text)

            # Step 6: Analyze cultural fit
            processing_steps.append("Analyzing cultural fit...")
            cultural_fit_scores = self.information_extractor.extract_cultural_fit(
                resume_text
            )

            # Step 7: Assess competencies
            processing_steps.append("Assessing competencies...")
            competency_assessments = self.competency_assessor.assess_competencies(
                resume_text, skills, work_experience
            )

            # Step 8: Comprehensive analysis
            processing_steps.append("Performing comprehensive analysis...")
            analysis_results = self._perform_comprehensive_analysis(
                resume_text,
                skills,
                work_experience,
                cultural_fit_scores,
                competency_assessments,
            )

            # Create analysis result
            analysis_id = (
                f"{personal_info.full_name}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
            )

            result = ResumeAnalysisResult(
                analysis_id=analysis_id,
                personal_info=personal_info,
                education=education,
                work_experience=work_experience,
                skills=skills,
                cultural_fit_scores=cultural_fit_scores,
                competency_assessments=competency_assessments,
                overall_score=analysis_results["overall_score"],
                early_talent_suitability=analysis_results["early_talent_suitability"],
                leadership_potential=analysis_results["leadership_potential"],
                career_progression_analysis=analysis_results[
                    "career_progression_analysis"
                ],
                strengths=analysis_results["strengths"],
                improvement_areas=analysis_results["improvement_areas"],
                red_flags=analysis_results["red_flags"],
                analysis_timestamp=datetime.now(),
                file_hash=file_hash,
                confidence_score=0.8,  # Will be assessed by G-Eval
                processing_metadata={
                    "processing_time": (datetime.now() - start_time).total_seconds(),
                    "file_path": str(file_path),
                    "text_length": len(resume_text),
                    "processing_steps": processing_steps,
                },
            )

            # Step 9: G-Eval quality assessment
            processing_steps.append("Evaluating extraction quality...")
            result.confidence_score = self.geval_assessor.evaluate_extraction_quality(
                resume_text, result
            )

            # Step 10: Save to database
            processing_steps.append("Saving analysis results...")
            save_success = self.database_manager.save_analysis_result(result)

            if not save_success:
                processing_steps.append("Warning: Failed to save to database")

            # Step 11: Input validation
            processing_steps.append("Validating extraction accuracy...")
            validation_results = self._validate_extraction(resume_text, result)

            processing_time = (datetime.now() - start_time).total_seconds()
            processing_steps.append(
                f"Processing completed in {processing_time:.2f} seconds"
            )

            # Format response
            response = {
                "answer": self._create_analysis_summary(result),
                "confidence": result.confidence_score,
                "sources": [f"Resume file: {Path(file_path).name}"],
                "workflow": "resume_reader",
                "analysis_id": result.analysis_id,
                "candidate_name": result.personal_info.full_name,
                "overall_score": result.overall_score,
                "early_talent_suitability": result.early_talent_suitability,
                "leadership_potential": result.leadership_potential,
                "extracted_data": {
                    "personal_info": self._serialize_personal_info(
                        result.personal_info
                    ),
                    "skills_count": len(result.skills),
                    "experience_count": len(result.work_experience),
                    "education_count": len(result.education),
                    "cultural_dimensions": len(result.cultural_fit_scores),
                    "competency_assessments": len(result.competency_assessments),
                },
                "validation_results": validation_results,
                "processing_steps": processing_steps,
                "timestamp": result.analysis_timestamp.isoformat(),
                "workflow_metadata": result.processing_metadata,
                "database_saved": save_success,
            }

            print(
                f"Resume analysis completed for {result.personal_info.full_name} with {result.confidence_score:.2f} confidence"
            )
            return response

        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            processing_steps.append(f"Error encountered: {str(e)}")

            return {
                "answer": f"Error encountered during resume analysis: {str(e)}",
                "confidence": 0.0,
                "sources": [],
                "workflow": "resume_reader_error",
                "error_details": {
                    "error_type": type(e).__name__,
                    "error_message": str(e),
                },
                "processing_steps": processing_steps,
                "timestamp": datetime.now().isoformat(),
                "workflow_metadata": {
                    "processing_time": processing_time,
                    "error_occurred": True,
                },
            }

    def handle_query(
        self, query: str, file_path: Optional[str] = None
    ) -> Dict[str, Any]:
        """Handle query (integration with existing system)"""

        if not file_path:
            return {
                "answer": "Please provide resume file path for analysis. Supported formats include: PDF, DOCX, TXT, and image files.",
                "confidence": 0.0,
                "sources": [],
                "workflow": "resume_reader",
                "timestamp": datetime.now().isoformat(),
            }

        # Process resume
        return self.process_resume(file_path)

    def _extract_work_experience_simple(self, resume_text: str) -> List[WorkExperience]:
        """Simplified work experience extraction"""
        # Basic keyword matching and pattern recognition
        # In practice, more sophisticated NLP techniques could be used

        experiences = []

        # Look for job-related keywords
        job_keywords = [
            "developer",
            "engineer",
            "manager",
            "analyst",
            "intern",
            "consultant",
        ]
        company_patterns = re.findall(
            r"(?:at|@)\s+([A-Za-z\s&.]{2,30})", resume_text, re.IGNORECASE
        )

        # Simplified extraction: create basic work experience entry
        if any(keyword in resume_text.lower() for keyword in job_keywords):
            experiences.append(
                WorkExperience(
                    position="Position information to be extracted",
                    company=(
                        company_patterns[0]
                        if company_patterns
                        else "Company information pending"
                    ),
                    start_date="To be confirmed",
                    end_date="To be confirmed",
                    responsibilities=[
                        "Responsibility information extracted from resume"
                    ],
                    achievements=["Achievement information extracted from resume"],
                )
            )

        return experiences

    def _extract_education_simple(self, resume_text: str) -> List[Education]:
        """Simplified education background extraction"""
        education = []

        # Look for education-related keywords
        education_keywords = [
            "university",
            "college",
            "bachelor",
            "master",
            "phd",
            "degree",
        ]

        if any(keyword in resume_text.lower() for keyword in education_keywords):
            education.append(
                Education(
                    degree="Degree information pending",
                    major="Major information pending",
                    school="School information pending",
                )
            )

        return education

    def _perform_comprehensive_analysis(
        self,
        resume_text: str,
        skills: List[Skill],
        work_experience: List[WorkExperience],
        cultural_fit_scores: List[CulturalFitScore],
        competency_assessments: List[CompetencyAssessment],
    ) -> Dict[str, Any]:
        """Perform comprehensive analysis"""

        # Calculate overall score
        skill_avg = (
            sum(skill.proficiency_score for skill in skills) / len(skills)
            if skills
            else 50
        )
        cultural_avg = (
            sum(score.score for score in cultural_fit_scores) / len(cultural_fit_scores)
            if cultural_fit_scores
            else 60
        )
        competency_avg = (
            sum(assessment.current_score for assessment in competency_assessments)
            / len(competency_assessments)
            if competency_assessments
            else 55
        )

        overall_score = skill_avg * 0.4 + cultural_avg * 0.3 + competency_avg * 0.3

        # Early talent suitability (based on experience years and learning ability)
        early_talent_suitability = min(
            85, overall_score + 10
        )  # Early talent usually has higher growth potential

        # Leadership potential assessment
        leadership_indicators = [
            assessment.current_score
            for assessment in competency_assessments
            if "leadership" in assessment.competency_name.lower()
            or "team" in assessment.competency_name.lower()
            or "communication" in assessment.competency_name.lower()
        ]
        leadership_potential = (
            sum(leadership_indicators) / len(leadership_indicators)
            if leadership_indicators
            else 60
        )

        # Identify strengths
        strengths = []
        for skill in skills:
            if skill.proficiency_score > 70:
                strengths.append(f"{skill.name} ({skill.category.value})")

        for assessment in competency_assessments:
            if assessment.current_score > 75:
                strengths.append(f"Strong {assessment.competency_name} capabilities")

        # Identify improvement areas
        improvement_areas = []
        for assessment in competency_assessments:
            if assessment.current_score < 50:
                improvement_areas.extend(assessment.development_areas)

        # Identify red flags
        red_flags = []
        if len(work_experience) == 0:
            red_flags.append("Lack of relevant work experience")

        if overall_score < 30:
            red_flags.append("Overall competency score is low")

        # Career progression analysis
        career_progression_analysis = f"""
        Based on resume analysis, the candidate's overall score is {overall_score:.1f}/100.
        Early talent suitability score is {early_talent_suitability:.1f}/100, showing {'good' if early_talent_suitability > 70 else 'moderate'} development potential.
        Leadership potential score is {leadership_potential:.1f}/100, {'showing advantages' if leadership_potential > 70 else 'needing strengthening'} in leadership development.
        Recommend focusing on {improvement_areas[:2] if improvement_areas else ['comprehensive capability enhancement']} for development.
        """

        return {
            "overall_score": overall_score,
            "early_talent_suitability": early_talent_suitability,
            "leadership_potential": leadership_potential,
            "career_progression_analysis": career_progression_analysis.strip(),
            "strengths": strengths[:5],  # Limit quantity
            "improvement_areas": improvement_areas[:5],
            "red_flags": red_flags,
        }

    def _create_analysis_summary(self, result: ResumeAnalysisResult) -> str:
        """Create analysis summary"""
        return f"""
        Resume Analysis Complete - {result.personal_info.full_name}
        
        === Basic Information ===
        Name: {result.personal_info.full_name}
        Contact: {result.personal_info.email or 'Not provided'} | {result.personal_info.phone or 'Not provided'}
        
        === Overall Scores ===
        Overall Score: {result.overall_score:.1f}/100
        Early Talent Suitability: {result.early_talent_suitability:.1f}/100  
        Leadership Potential: {result.leadership_potential:.1f}/100
        
        === Skills Overview ===
        Skills Identified: {len(result.skills)} items
        Cultural Fit Dimensions: {len(result.cultural_fit_scores)} dimensions
        Competency Assessments: {len(result.competency_assessments)} items
        
        === Key Strengths ===
        {chr(10).join(['• ' + strength for strength in result.strengths[:3]])}
        
        === Development Suggestions ===
        {chr(10).join(['• ' + area for area in result.improvement_areas[:3]])}
        
        Analysis Confidence: {result.confidence_score:.2f}
        Data saved to database with Analysis ID: {result.analysis_id}
        """

    def _serialize_personal_info(self, info: PersonalInfo) -> Dict[str, Any]:
        """Serialize personal information"""
        return {
            "full_name": info.full_name,
            "email": info.email,
            "phone": info.phone,
            "address": info.address,
            "linkedin": info.linkedin,
            "github": info.github,
            "portfolio": info.portfolio,
            "summary": info.summary,
        }

    def _validate_extraction(
        self, original_text: str, result: ResumeAnalysisResult
    ) -> Dict[str, Any]:
        """Validate extraction accuracy"""
        validation = {
            "basic_info_present": bool(result.personal_info.full_name != "Unknown"),
            "email_extracted": bool(result.personal_info.email),
            "phone_extracted": bool(result.personal_info.phone),
            "skills_extracted": len(result.skills) > 0,
            "text_length_reasonable": len(original_text) > 100,
            "overall_scores_reasonable": 0 <= result.overall_score <= 100,
        }

        validation["validation_score"] = sum(validation.values()) / len(validation)
        return validation

    def _create_initialization_error_response(self) -> Dict[str, Any]:
        """Create initialization error response"""
        return {
            "answer": "Resume reader system initialization failed, please contact technical support.",
            "confidence": 0.0,
            "sources": [],
            "workflow": "resume_reader_init_error",
            "timestamp": datetime.now().isoformat(),
            "error_details": {
                "initialization_error": getattr(
                    self, "initialization_error", "Unknown error"
                )
            },
        }

    def _create_error_response(
        self, error_message: str, processing_steps: List[str]
    ) -> Dict[str, Any]:
        """Create error response"""
        return {
            "answer": f"Problem encountered during resume analysis: {error_message}",
            "confidence": 0.0,
            "sources": [],
            "workflow": "resume_reader_error",
            "processing_steps": processing_steps,
            "timestamp": datetime.now().isoformat(),
            "error_details": {"error_message": error_message},
        }

    def get_candidate_history(self, candidate_name: str) -> Dict[str, Any]:
        """Get candidate history analysis records"""
        history = self.database_manager.get_candidate_history(candidate_name)

        if not history:
            return {
                "candidate_name": candidate_name,
                "history_found": False,
                "message": "No historical analysis records found for this candidate",
            }

        # Calculate improvement trend
        if len(history) > 1:
            latest = history[0]
            previous = history[1]

            score_change = latest["overall_score"] - previous["overall_score"]
            trend = (
                "Improving"
                if score_change > 0
                else "Declining" if score_change < 0 else "Stable"
            )
        else:
            trend = "First analysis"

        return {
            "candidate_name": candidate_name,
            "history_found": True,
            "analysis_count": len(history),
            "latest_analysis": history[0] if history else None,
            "trend": trend,
            "full_history": history,
        }

    def get_system_status(self) -> Dict[str, Any]:
        """Get system status"""
        return {
            "initialization_successful": self.initialization_successful,
            "initialization_time": self.initialization_time.isoformat(),
            "langchain_available": LANGCHAIN_AVAILABLE,
            "pdf_processing": PDF_AVAILABLE,
            "docx_processing": DOCX_AVAILABLE,
            "ocr_available": OCR_AVAILABLE,
            "deepeval_available": DEEPEVAL_AVAILABLE,
            "supported_formats": self.config.supported_formats,
            "database_path": str(self.config.database_path),
            "config": {
                "enable_ocr": self.config.enable_ocr,
                "enable_cultural_analysis": self.config.enable_cultural_analysis,
                "min_confidence_threshold": self.config.min_confidence_threshold,
            },
        }


# ========================================================================================
# UTILITY FUNCTIONS
# ========================================================================================


def create_sample_resume_text() -> str:
    """Create sample resume text for testing"""
    return """
    John Smith
    Email: john.smith@email.com
    Phone: +1-555-0123
    LinkedIn: linkedin.com/in/johnsmith
    
    Career Objective
    Seeking a software development engineer position to utilize technical innovation and teamwork skills.
    
    Education
    Bachelor's Degree, Computer Science
    University of California, Berkeley, 2020-2024
    GPA: 3.7/4.0
    
    Work Experience
    Software Development Intern
    Google Inc., June 2023 - December 2023
    • Participated in Android app development projects
    • Used JavaScript and Python for front-end and back-end development
    • Collaborated with team to complete product feature optimization
    
    Skills
    • Programming Languages: Python (Advanced), JavaScript (Intermediate), Java (Intermediate)
    • Frameworks: React, Django, Spring Boot
    • Databases: MySQL, MongoDB
    • Tools: Git, Docker, Jenkins
    • Soft Skills: Teamwork, Communication, Problem Solving, Strong Learning Ability
    
    Project Experience
    Online Learning Platform
    • Full-stack web application developed using React and Django
    • Implemented user authentication, course management, and online testing features
    • Deployed on AWS cloud platform, supporting 1000+ concurrent users
    
    Certifications and Honors
    • AWS Cloud Practitioner Certification
    • Second Prize in University Programming Competition
    • Outstanding Student Scholarship Recipient
    """


# ========================================================================================
# MAIN EXECUTION (FOR TESTING)
# ========================================================================================

if __name__ == "__main__":
    print("Resume Reader Component Demo")
    print("=" * 60)

    # Initialize mock components
    llm = MockLLM()

    # Create configuration
    config = ReaderConfig(
        database_path="test_resumes.db",
        enable_ocr=False,  # Disable OCR in testing
        enable_cultural_analysis=True,
    )

    # Initialize workflow
    workflow = ReaderWorkflow(llm, config)

    # Create test resume file
    test_resume_path = Path("test_resume.txt")
    test_resume_path.write_text(create_sample_resume_text(), encoding="utf-8")

    print(f"\nTesting Resume Analysis:")
    print("-" * 60)

    # Test resume processing
    try:
        result = workflow.process_resume(str(test_resume_path), "John Smith")

        print(f"Analysis Result:")
        print(f"Candidate: {result['candidate_name']}")
        print(f"Overall Score: {result['overall_score']:.1f}/100")
        print(f"Early Talent Suitability: {result['early_talent_suitability']:.1f}/100")
        print(f"Leadership Potential: {result['leadership_potential']:.1f}/100")
        print(f"Confidence: {result['confidence']:.2f}")
        print(f"Database Saved: {result['database_saved']}")

        print(f"\nExtracted Data Summary:")
        extracted = result["extracted_data"]
        print(f"  Skills: {extracted['skills_count']}")
        print(f"  Experience: {extracted['experience_count']}")
        print(f"  Education: {extracted['education_count']}")
        print(f"  Cultural Dimensions: {extracted['cultural_dimensions']}")
        print(f"  Competency Assessments: {extracted['competency_assessments']}")

    except Exception as e:
        print(f"Error during testing: {e}")

    # Test candidate history query
    print(f"\n" + "=" * 60)
    print("Testing Candidate History:")
    print("-" * 60)

    history_result = workflow.get_candidate_history("John Smith")
    print(f"History Found: {history_result['history_found']}")
    if history_result["history_found"]:
        print(f"Analysis Count: {history_result['analysis_count']}")
        print(f"Trend: {history_result['trend']}")

    # Display system status
    print(f"\n" + "=" * 60)
    print("System Status:")
    print("-" * 60)
    status = workflow.get_system_status()
    print(
        f"Initialization: {'Success' if status['initialization_successful'] else 'Failed'}"
    )
    print(f"PDF Processing: {status['pdf_processing']}")
    print(f"DOCX Processing: {status['docx_processing']}")
    print(f"OCR Available: {status['ocr_available']}")
    print(f"Supported Formats: {', '.join(status['supported_formats'])}")

    # Clean up test file
    if test_resume_path.exists():
        test_resume_path.unlink()

    print(f"\nResume Reader Component Demo Completed!")
